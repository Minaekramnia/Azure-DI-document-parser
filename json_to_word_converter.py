"""
Convert JSON Analysis Files to Word Documents

This script converts the JSON output files (*_WORKING_analysis.json) 
to formatted Word documents for review and sharing.
"""

import json
import glob
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def convert_json_to_word(json_filepath):
    """
    Convert a single JSON analysis file to Word document.
    
    Args:
        json_filepath: Path to the JSON file
    
    Returns:
        Document: python-docx Document object
    """
    try:
        with open(json_filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        doc = Document()
        doc.add_heading('Error Loading File', 0)
        doc.add_paragraph(f"Could not load {json_filepath}: {e}")
        return doc
    
    # Create Word document
    doc = Document()
    
    # Title
    doc_name = os.path.basename(data.get('document_path', 'Unknown'))
    title = doc.add_heading(f'Document Analysis: {doc_name}', 0)
    
    # Metadata
    doc.add_paragraph(f"Source: {data.get('document_path', 'Unknown')}")
    doc.add_paragraph(f"Total Pages: {data.get('total_pages', 0)}")
    doc.add_paragraph(f"Total Segments: {data.get('total_segments', 0)}")
    doc.add_paragraph(f"Analysis: Executive Prompt (AI Document Analysis)")
    doc.add_paragraph()
    
    # Summary section
    segments = data.get('segments', [])
    all_names = set()
    all_categories = set()
    total_classifications = 0
    
    for segment in segments:
        all_names.update(segment.get('extracted_names', []))
        for classification in segment.get('classifications', []):
            all_categories.add(classification.get('category', 'Unknown'))
            total_classifications += 1
    
    summary_heading = doc.add_heading('Summary', 1)
    doc.add_paragraph(f"Total Names Extracted: {len(all_names)}")
    doc.add_paragraph(f"Total Sensitive Items Flagged: {total_classifications}")
    doc.add_paragraph(f"Categories Found: {', '.join(sorted(all_categories)) if all_categories else 'None'}")
    doc.add_paragraph()
    
    if all_names:
        doc.add_heading('All Names Found:', 2)
        for name in sorted(all_names):
            p = doc.add_paragraph(name, style='List Bullet')
        doc.add_paragraph()
    
    # Process each segment
    for segment in segments:
        segment_id = segment.get('segment_id', 'Unknown')
        page_range = segment.get('page_range', 'Unknown')
        
        # Segment heading
        doc.add_heading(f"{segment_id.replace('_', ' ').title()}: Pages {page_range}", 1)
        
        # Extracted names
        names = segment.get('extracted_names', [])
        doc.add_heading('Extracted Names', 2)
        if names:
            for name in names:
                p = doc.add_paragraph(name, style='List Bullet')
                p.runs[0].bold = True
        else:
            p = doc.add_paragraph('No names found in this segment')
            p.italic = True
        doc.add_paragraph()
        
        # Classifications
        classifications = segment.get('classifications', [])
        doc.add_heading('Sensitive Content Classifications', 2)
        
        if classifications:
            for i, classification in enumerate(classifications, 1):
                category = classification.get('category', 'Unknown')
                text = classification.get('text', '')
                bbox = classification.get('bounding_box', [])
                confidence = classification.get('confidence_score', 0.0)
                reason = classification.get('reason', '')
                
                # Classification heading
                class_heading = doc.add_heading(f"Classification {i}: {category}", 3)
                
                # Text
                doc.add_paragraph('Text:', style='Heading 4')
                text_para = doc.add_paragraph(text[:500] + ('...' if len(text) > 500 else ''))
                text_para.runs[0].font.color.rgb = RGBColor(50, 50, 50)
                
                # Confidence
                confidence_para = doc.add_paragraph(f"Confidence Score: {confidence:.2f} ({int(confidence * 100)}%)")
                confidence_para.runs[0].bold = True
                
                # Reason
                doc.add_paragraph(f"Reason: {reason}")
                
                # Bounding box
                if bbox:
                    bbox_str = f"[{', '.join(map(str, bbox))}]"
                    doc.add_paragraph(f"Bounding Box: {bbox_str}")
                
                doc.add_paragraph()
        else:
            p = doc.add_paragraph('No sensitive content found in this segment')
            p.italic = True
        
        doc.add_paragraph()
    
    # Footer
    doc.add_page_break()
    footer_heading = doc.add_heading('Document Processing Complete', 1)
    doc.add_paragraph('This document has been analyzed using the Executive Prompt for sensitive information detection.')
    
    return doc


def convert_all_json_files(data_path, output_folder_name='word_reports'):
    """
    Convert all JSON analysis files to Word documents.
    
    Args:
        data_path: Directory containing the JSON files
        output_folder_name: Name of the output folder to create
    """
    # Find all JSON analysis files
    json_files = glob.glob(f"{data_path}*_WORKING_analysis.json")
    
    if not json_files:
        print(f"❌ No *_WORKING_analysis.json files found in {data_path}")
        return
    
    # Create output folder
    output_path = os.path.join(data_path, output_folder_name)
    try:
        os.makedirs(output_path, exist_ok=True)
        print(f"✅ Output folder created: {output_path}")
    except Exception as e:
        print(f"❌ Could not create output folder: {e}")
        return
    
    print("\n🔄 JSON to Word Converter")
    print("="*60)
    print(f"Input: {data_path}")
    print(f"Output: {output_path}/")
    print(f"Found {len(json_files)} JSON files to convert")
    print("="*60)
    
    converted_count = 0
    error_count = 0
    
    for json_file in sorted(json_files):
        print(f"\n📄 Converting: {os.path.basename(json_file)}")
        
        try:
            # Convert to Word
            word_doc = convert_json_to_word(json_file)
            
            # Create Word filename
            base_name_docx = os.path.basename(json_file).replace('_WORKING_analysis.json', '_WORKING_analysis.docx')
            docx_filename = os.path.join(output_path, base_name_docx)
            
            # Save Word document
            word_doc.save(docx_filename)
            
            print(f"   ✅ Created: {base_name_docx}")
            converted_count += 1
            
        except Exception as e:
            print(f"   ❌ Error: {e}")
            error_count += 1
    
    print("\n" + "="*60)
    print("🎯 Conversion Complete!")
    print("="*60)
    print(f"✅ Successfully converted: {converted_count} files")
    if error_count > 0:
        print(f"❌ Errors: {error_count} files")
    print(f"\n📁 Output folder: {output_path}/")
    print(f"📄 Word files: *_WORKING_analysis.docx")
    print("="*60)


if __name__ == '__main__':
    # Path to the directory containing JSON files
    data_path = '/Volumes/qa_datascience_convaiqa/volumes/dataanalyticslake_convai/PRISMArchives/PI/'
    
    print("\n" + "="*60)
    print("🚀 JSON TO WORD CONVERTER")
    print("="*60)
    print(f"Source: {data_path}")
    print(f"Looking for: *_WORKING_analysis.json")
    print(f"Output: word_reports/ folder")
    print("="*60)
    
    convert_all_json_files(data_path)
    
    print("\n✅ All JSON files converted to Word documents!")
    print("   📄 Open the .docx files in Microsoft Word for review and editing.")


