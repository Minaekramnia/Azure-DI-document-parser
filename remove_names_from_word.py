"""
Remove Extracted Names from Word Documents

This script:
1. Finds all Word documents in word_reports/ folder
2. Removes the "Extracted Names" sections from each document
3. Saves cleaned documents to final_word_reports/ folder

Output Path:
- Final Word: /Volumes/.../PI/final_word_reports/
"""

import glob
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("❌ python-docx not installed. Run: pip install python-docx")
    exit(1)


def remove_names_from_word(input_filepath, output_filepath):
    """
    Remove extracted names sections from a Word document.
    
    Args:
        input_filepath: Path to input Word file
        output_filepath: Path for output Word file
    
    Returns:
        bool: Success status
    """
    try:
        # Load the document
        doc = Document(input_filepath)
        
        # Find and remove paragraphs containing "Extracted Names"
        paragraphs_to_remove = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Check if this is an "Extracted Names" section
            if text == "Extracted Names":
                # Mark this paragraph and the next few for removal
                paragraphs_to_remove.append(i)
                
                # Find the end of the extracted names section
                # (look for next section header or end of document)
                for j in range(i + 1, len(doc.paragraphs)):
                    next_text = doc.paragraphs[j].text.strip()
                    
                    # Stop at next section headers
                    if (next_text in ["Sensitive Content Classifications", "Summary Statistics", 
                                     "Classification Breakdown", "Segment:", ""] or
                        next_text.startswith("Classification") or
                        next_text.startswith("Segment:")):
                        break
                    
                    # If it's a name (starts with "- "), remove it
                    if next_text.startswith("- "):
                        paragraphs_to_remove.append(j)
                    # If it's "No names extracted", remove it
                    elif next_text == "No names extracted":
                        paragraphs_to_remove.append(j)
                    # If it's empty line after names, remove it
                    elif next_text == "" and j < len(doc.paragraphs) - 1:
                        next_next_text = doc.paragraphs[j + 1].text.strip()
                        if (next_next_text in ["Sensitive Content Classifications", "Summary Statistics"] or
                            next_next_text.startswith("Classification") or
                            next_next_text.startswith("Segment:")):
                            paragraphs_to_remove.append(j)
                            break
        
        # Remove paragraphs in reverse order (to maintain indices)
        for i in sorted(paragraphs_to_remove, reverse=True):
            if i < len(doc.paragraphs):
                # Remove the paragraph
                p = doc.paragraphs[i]
                p_element = p._element
                p_element.getparent().remove(p_element)
        
        # Save the cleaned document
        doc.save(output_filepath)
        return True
        
    except Exception as e:
        print(f"   ❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Main processing function."""
    
    # Paths
    data_path = '/Volumes/qa_datascience_convaiqa/volumes/dataanalyticslake_convai/PRISMArchives/PI/'
    word_folder = os.path.join(data_path, 'word_reports')
    final_folder = os.path.join(data_path, 'final_word_reports')
    
    # Find Word files
    word_files = glob.glob(f"{word_folder}/*.docx")
    
    if not word_files:
        print(f"❌ No Word files found in {word_folder}")
        print(f"   Looking for: *.docx files")
        return
    
    # Create final_word_reports folder
    try:
        os.makedirs(final_folder, exist_ok=True)
        print(f"✅ Final folder: {final_folder}\n")
    except Exception as e:
        print(f"❌ Could not create folder: {e}")
        return
    
    print("="*80)
    print("🚀 REMOVING EXTRACTED NAMES FROM WORD DOCUMENTS")
    print("="*80)
    print(f"Input: {word_folder}")
    print(f"Output: {final_folder}")
    print(f"Found {len(word_files)} Word files to process")
    print("="*80 + "\n")
    
    success_count = 0
    error_count = 0
    
    for word_file in sorted(word_files):
        filename = os.path.basename(word_file)
        print(f"📄 {filename}")
        
        # Create output filename
        final_file = os.path.join(final_folder, filename)
        
        # Process file
        print(f"   🗑️  Removing extracted names...")
        if remove_names_from_word(word_file, final_file):
            print(f"   ✅ Success: {filename}")
            success_count += 1
        else:
            print(f"   ❌ Failed")
            error_count += 1
        
        print()
    
    # Summary
    print("="*80)
    print("🎯 PROCESSING COMPLETE!")
    print("="*80)
    print(f"✅ Files processed: {success_count}")
    if error_count > 0:
        print(f"❌ Errors: {error_count}")
    print()
    print(f"📁 Original files: {word_folder}")
    print(f"📁 Final files: {final_folder}")
    print("="*80)
    
    print(f"\n📍 OUTPUT PATHS:")
    print(f"   📄 Original (with names): {word_folder}")
    print(f"   📄 Final (names removed): {final_folder}\n")
    
    print("💡 Note: Original files are preserved in word_reports/")
    print("   Final cleaned files are in final_word_reports/")


if __name__ == '__main__':
    if not DOCX_AVAILABLE:
        print("\n❌ Cannot proceed without python-docx")
        print("   Install with: pip install python-docx\n")
    else:
        main()
