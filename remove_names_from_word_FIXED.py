"""
Remove Extracted Names from Word Documents - FIXED for Databricks Volumes

This script:
1. Finds all Word documents in word_reports/ folder
2. Removes the "Extracted Names" sections from each document
3. Saves cleaned documents to final_word_reports/ folder
4. Uses temp file workaround for Databricks Volumes

Output Path:
- Final Word: /Volumes/.../PI/final_word_reports/
"""

import glob
import os
import shutil
import tempfile
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("❌ python-docx not installed. Run: pip install python-docx")
    exit(1)


def remove_names_from_word(input_filepath, output_filepath):
    """
    Remove extracted names sections from a Word document with filesystem workaround.
    
    Args:
        input_filepath: Path to input Word file
        output_filepath: Path for output Word file
    
    Returns:
        bool: Success status
    """
    try:
        # Load the document
        doc = Document(input_filepath)
        
        # Find paragraphs to remove
        paragraphs_to_remove = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Check if this is an "Extracted Names" section
            if text == "Extracted Names":
                paragraphs_to_remove.append(i)
                
                # Find related paragraphs to remove
                for j in range(i + 1, len(doc.paragraphs)):
                    next_text = doc.paragraphs[j].text.strip()
                    
                    # Stop at next section headers
                    if (next_text in ["Sensitive Content Classifications", "Summary Statistics", 
                                     "Classification Breakdown", "Segment:", ""] or
                        next_text.startswith("Classification") or
                        next_text.startswith("Segment:")):
                        break
                    
                    # Remove name items and empty lines
                    if (next_text.startswith("- ") or 
                        next_text == "No names extracted" or
                        (next_text == "" and j < len(doc.paragraphs) - 1)):
                        paragraphs_to_remove.append(j)
        
        # Remove paragraphs in reverse order (to maintain indices)
        for i in sorted(paragraphs_to_remove, reverse=True):
            if i < len(doc.paragraphs):
                p = doc.paragraphs[i]
                p_element = p._element
                p_element.getparent().remove(p_element)
        
        # WORKAROUND: Save to temp file first, then copy
        # This avoids "Operation not supported" error on Databricks Volumes
        with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as tmp_file:
            temp_path = tmp_file.name
        
        # Save to temp location
        doc.save(temp_path)
        
        # Copy to final destination
        shutil.copy2(temp_path, output_filepath)
        
        # Clean up temp file
        try:
            os.remove(temp_path)
        except:
            pass
        
        return True
        
    except Exception as e:
        print(f"   ❌ Error: {e}")
        return False


def main():
    """Main processing function."""
    
    # Paths
    data_path = '/Volumes/qa_datascience_convaiqa/volumes/dataanalyticslake_convai/PRISMArchives/PI/'
    word_folder = os.path.join(data_path, 'word_reports')
    final_folder = os.path.join(data_path, 'final_word_reports')
    
    # Check if input folder exists
    if not os.path.exists(word_folder):
        print(f"❌ Input folder does not exist: {word_folder}")
        print("   Please run convert_to_word_fixed.py first to create Word documents")
        return
    
    # Find Word files
    word_files = glob.glob(f"{word_folder}/*.docx")
    
    if not word_files:
        print(f"❌ No Word files found in {word_folder}")
        print(f"   Looking for: *.docx files")
        print(f"   Folder contents: {os.listdir(word_folder) if os.path.exists(word_folder) else 'Folder not found'}")
        return
    
    # Create final_word_reports folder
    try:
        os.makedirs(final_folder, exist_ok=True)
        print(f"✅ Final folder created: {final_folder}")
    except Exception as e:
        print(f"❌ Could not create folder: {e}")
        return
    
    print("\n" + "="*80)
    print("🚀 REMOVING EXTRACTED NAMES FROM WORD DOCUMENTS")
    print("="*80)
    print(f"Input: {word_folder}")
    print(f"Output: {final_folder}")
    print(f"Found {len(word_files)} Word files to process")
    print("="*80 + "\n")
    
    success_count = 0
    error_count = 0
    error_details = []
    
    for word_file in sorted(word_files):
        filename = os.path.basename(word_file)
        print(f"📄 Processing: {filename}")
        
        try:
            # Create output filename
            final_file = os.path.join(final_folder, filename)
            
            # Process file
            print(f"   🗑️  Removing extracted names...")
            if remove_names_from_word(word_file, final_file):
                print(f"   ✅ Success: {filename}")
                success_count += 1
            else:
                print(f"   ❌ Failed: {filename}")
                error_count += 1
                error_details.append(filename)
        
        except Exception as e:
            print(f"   ❌ Exception: {filename} - {e}")
            error_count += 1
            error_details.append(f"{filename}: {e}")
        
        print()
    
    # Summary
    print("="*80)
    print("🎯 PROCESSING COMPLETE!")
    print("="*80)
    print(f"✅ Files processed successfully: {success_count}")
    if error_count > 0:
        print(f"❌ Files with errors: {error_count}")
        print("Error details:")
        for detail in error_details:
            print(f"   - {detail}")
    print()
    print(f"📁 Original files: {word_folder}")
    print(f"📁 Final files: {final_folder}")
    print("="*80)
    
    print(f"\n📍 OUTPUT PATHS:")
    print(f"   📄 Original (with names): {word_folder}")
    print(f"   📄 Final (names removed): {final_folder}\n")
    
    if success_count > 0:
        print("💡 Note: Original files are preserved in word_reports/")
        print("   Final cleaned files are in final_word_reports/")
    else:
        print("⚠️  No files were processed successfully. Check error details above.")


if __name__ == '__main__':
    if not DOCX_AVAILABLE:
        print("\n❌ Cannot proceed without python-docx")
        print("   Install with: pip install python-docx\n")
    else:
        main()
