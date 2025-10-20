"""
Convert PageNumbers JSON Analysis to Markdown and Word Documents

This script:
1. Finds all *_WORKING_PageNumbers_analysis.json files
2. Converts them to Markdown format (with page number tracking)
3. Converts to Word documents using python-docx (no pandoc required!)
4. Saves outputs in organized folders

Output Paths:
- Markdown: /Volumes/.../PI/markdown_reports/
- Word: /Volumes/.../PI/word_reports/
"""

import json
import glob
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not available. Install with: pip install python-docx")


def convert_json_to_markdown(json_filepath, output_filepath):
    """
    Convert JSON analysis to human-readable Markdown format.
    Includes page number information for each classification.
    
    Args:
        json_filepath: Path to the JSON file
        output_filepath: Path for the output .md file
    
    Returns:
        tuple: (success: bool, data: dict or None)
    """
    try:
        # Load JSON
        with open(json_filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Start building Markdown
        md_content = []
        
        # Header
        md_content.append("# Document Analysis Report (with Page Numbers)\n")
        md_content.append(f"**Generated**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        md_content.append(f"**Source Document**: `{os.path.basename(data.get('document_path', 'Unknown'))}`\n")
        md_content.append(f"**Total Pages**: {data.get('total_pages', 'Unknown')}\n")
        md_content.append(f"**Total Segments**: {data.get('total_segments', 'Unknown')}\n")
        md_content.append("\n---\n\n")
        
        # Process each segment
        segments = data.get('segments', [])
        
        if not segments:
            md_content.append("*No segments found in this document.*\n")
        
        for segment in segments:
            segment_id = segment.get('segment_id', 'Unknown')
            page_range = segment.get('page_range', 'Unknown')
            
            md_content.append(f"## Segment: {segment_id}\n")
            md_content.append(f"**Pages**: {page_range}\n\n")
            
            # Extracted Names
            extracted_names = segment.get('extracted_names', [])
            md_content.append("### 📋 Extracted Names\n")
            if extracted_names:
                for name in extracted_names:
                    md_content.append(f"- {name}\n")
            else:
                md_content.append("*No names extracted*\n")
            md_content.append("\n")
            
            # Classifications
            classifications = segment.get('classifications', [])
            md_content.append("### 🔍 Sensitive Content Classifications\n")
            
            if not classifications:
                md_content.append("*No sensitive content identified*\n\n")
            else:
                md_content.append(f"**Total Classifications**: {len(classifications)}\n\n")
                
                for i, classification in enumerate(classifications, 1):
                    category = classification.get('category', 'Unknown')
                    text = classification.get('text', 'N/A')
                    page_number = classification.get('page_number', 'N/A')
                    page_range_cv = classification.get('page_range', None)
                    confidence = classification.get('confidence_score', 'N/A')
                    reason = classification.get('reason', 'N/A')
                    bounding_box = classification.get('bounding_box', [])
                    
                    md_content.append(f"#### Classification {i}\n")
                    md_content.append(f"**Category**: `{category}`\n\n")
                    
                    # Page information (with special handling for CVs)
                    if page_range_cv:
                        md_content.append(f"**Page**: {page_number} (Full range: {page_range_cv})\n\n")
                    else:
                        md_content.append(f"**Page**: {page_number}\n\n")
                    
                    md_content.append(f"**Confidence Score**: {confidence}\n\n")
                    
                    md_content.append(f"**Content**:\n")
                    md_content.append(f"> {text}\n\n")
                    
                    md_content.append(f"**Reason**:\n")
                    md_content.append(f"> {reason}\n\n")
                    
                    if bounding_box:
                        bbox_str = ", ".join(map(str, bounding_box))
                        md_content.append(f"**Bounding Box**: `[{bbox_str}]`\n\n")
                    
                    md_content.append("---\n\n")
            
            md_content.append("\n")
        
        # Summary Statistics
        md_content.append("## 📊 Summary Statistics\n\n")
        
        total_names = sum(len(seg.get('extracted_names', [])) for seg in segments)
        total_classifications = sum(len(seg.get('classifications', [])) for seg in segments)
        
        md_content.append(f"- **Total Names Extracted**: {total_names}\n")
        md_content.append(f"- **Total Classifications**: {total_classifications}\n")
        md_content.append(f"- **Total Segments**: {len(segments)}\n")
        md_content.append(f"- **Total Pages**: {data.get('total_pages', 'Unknown')}\n\n")
        
        # Category breakdown
        category_counts = {}
        for segment in segments:
            for classification in segment.get('classifications', []):
                category = classification.get('category', 'Unknown')
                category_counts[category] = category_counts.get(category, 0) + 1
        
        if category_counts:
            md_content.append("### Classification Breakdown by Category\n\n")
            for category, count in sorted(category_counts.items(), key=lambda x: x[1], reverse=True):
                md_content.append(f"- **{category}**: {count}\n")
        
        md_content.append("\n---\n\n")
        md_content.append("*Report generated by Azure DI Output Parser (PageNumbers Version)*\n")
        
        # Write to file
        with open(output_filepath, 'w', encoding='utf-8') as f:
            f.write(''.join(md_content))
        
        return True, data
        
    except Exception as e:
        print(f"   ❌ Error converting to Markdown: {e}")
        return False, None


def convert_json_to_word(json_data, output_filepath):
    """
    Convert JSON data directly to Word document using python-docx.
    
    Args:
        json_data: Parsed JSON data
        output_filepath: Path for the output .docx file
    
    Returns:
        bool: True if successful, False otherwise
    """
    if not DOCX_AVAILABLE:
        print(f"   ⚠️  python-docx not installed. Skipping Word conversion.")
        return False
    
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Document Analysis Report (with Page Numbers)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Source Document: {os.path.basename(json_data.get('document_path', 'Unknown'))}")
        doc.add_paragraph(f"Total Pages: {json_data.get('total_pages', 'Unknown')}")
        doc.add_paragraph(f"Total Segments: {json_data.get('total_segments', 'Unknown')}")
        doc.add_paragraph()
        
        # Process each segment
        segments = json_data.get('segments', [])
        
        if not segments:
            doc.add_paragraph("No segments found in this document.", style='Intense Quote')
        
        for segment in segments:
            segment_id = segment.get('segment_id', 'Unknown')
            page_range = segment.get('page_range', 'Unknown')
            
            # Segment header
            doc.add_heading(f'Segment: {segment_id}', 1)
            doc.add_paragraph(f"Pages: {page_range}")
            
            # Extracted Names
            doc.add_heading('📋 Extracted Names', 2)
            extracted_names = segment.get('extracted_names', [])
            if extracted_names:
                for name in extracted_names:
                    doc.add_paragraph(name, style='List Bullet')
            else:
                doc.add_paragraph("No names extracted", style='Intense Quote')
            
            # Classifications
            doc.add_heading('🔍 Sensitive Content Classifications', 2)
            classifications = segment.get('classifications', [])
            
            if not classifications:
                doc.add_paragraph("No sensitive content identified", style='Intense Quote')
            else:
                doc.add_paragraph(f"Total Classifications: {len(classifications)}")
                doc.add_paragraph()
                
                for i, classification in enumerate(classifications, 1):
                    category = classification.get('category', 'Unknown')
                    text = classification.get('text', 'N/A')
                    page_number = classification.get('page_number', 'N/A')
                    page_range_cv = classification.get('page_range', None)
                    confidence = classification.get('confidence_score', 'N/A')
                    reason = classification.get('reason', 'N/A')
                    bounding_box = classification.get('bounding_box', [])
                    
                    # Classification header
                    doc.add_heading(f'Classification {i}', 3)
                    
                    # Category
                    p = doc.add_paragraph()
                    p.add_run('Category: ').bold = True
                    p.add_run(category)
                    
                    # Page information
                    p = doc.add_paragraph()
                    p.add_run('Page: ').bold = True
                    if page_range_cv:
                        p.add_run(f"{page_number} (Full range: {page_range_cv})")
                    else:
                        p.add_run(str(page_number))
                    
                    # Confidence
                    p = doc.add_paragraph()
                    p.add_run('Confidence Score: ').bold = True
                    p.add_run(str(confidence))
                    
                    # Content
                    p = doc.add_paragraph()
                    p.add_run('Content:').bold = True
                    doc.add_paragraph(text, style='Intense Quote')
                    
                    # Reason
                    p = doc.add_paragraph()
                    p.add_run('Reason:').bold = True
                    doc.add_paragraph(reason, style='Intense Quote')
                    
                    # Bounding box
                    if bounding_box:
                        bbox_str = ", ".join(map(str, bounding_box))
                        p = doc.add_paragraph()
                        p.add_run('Bounding Box: ').bold = True
                        p.add_run(f"[{bbox_str}]")
                    
                    doc.add_paragraph()  # Spacing
        
        # Summary Statistics
        doc.add_page_break()
        doc.add_heading('📊 Summary Statistics', 1)
        
        total_names = sum(len(seg.get('extracted_names', [])) for seg in segments)
        total_classifications = sum(len(seg.get('classifications', [])) for seg in segments)
        
        doc.add_paragraph(f"Total Names Extracted: {total_names}", style='List Bullet')
        doc.add_paragraph(f"Total Classifications: {total_classifications}", style='List Bullet')
        doc.add_paragraph(f"Total Segments: {len(segments)}", style='List Bullet')
        doc.add_paragraph(f"Total Pages: {json_data.get('total_pages', 'Unknown')}", style='List Bullet')
        
        # Category breakdown
        category_counts = {}
        for segment in segments:
            for classification in segment.get('classifications', []):
                category = classification.get('category', 'Unknown')
                category_counts[category] = category_counts.get(category, 0) + 1
        
        if category_counts:
            doc.add_heading('Classification Breakdown by Category', 2)
            for category, count in sorted(category_counts.items(), key=lambda x: x[1], reverse=True):
                doc.add_paragraph(f"{category}: {count}", style='List Bullet')
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph("Report generated by Azure DI Output Parser (PageNumbers Version)")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save
        doc.save(output_filepath)
        return True
        
    except Exception as e:
        print(f"   ❌ Error converting to Word: {e}")
        return False


def process_all_files(data_path):
    """
    Process all PageNumbers JSON files to create Markdown and Word reports.
    
    Args:
        data_path: Directory containing the JSON files
    """
    # Find all PageNumbers JSON files
    json_files = glob.glob(f"{data_path}*_WORKING_PageNumbers_analysis.json")
    
    if not json_files:
        print(f"❌ No *_WORKING_PageNumbers_analysis.json files found in {data_path}")
        return
    
    # Create output folders (simplified names)
    markdown_folder = os.path.join(data_path, 'markdown_reports')
    word_folder = os.path.join(data_path, 'word_reports')
    
    try:
        os.makedirs(markdown_folder, exist_ok=True)
        print(f"✅ Markdown folder created: {markdown_folder}")
    except Exception as e:
        print(f"❌ Could not create markdown folder: {e}")
        return
    
    try:
        os.makedirs(word_folder, exist_ok=True)
        print(f"✅ Word folder created: {word_folder}")
    except Exception as e:
        print(f"❌ Could not create word folder: {e}")
        return
    
    print("\n" + "="*80)
    print("🚀 CONVERTING PAGENUMBERS JSON TO MARKDOWN AND WORD")
    print("="*80)
    print(f"Input: {data_path}")
    print(f"Markdown Output: {markdown_folder}")
    print(f"Word Output: {word_folder}")
    print(f"Found {len(json_files)} JSON files to convert")
    print("="*80 + "\n")
    
    markdown_success = 0
    word_success = 0
    errors = 0
    
    for json_file in sorted(json_files):
        print(f"📄 Processing: {os.path.basename(json_file)}")
        
        try:
            # Create output filenames
            base_name = os.path.basename(json_file).replace('_WORKING_PageNumbers_analysis.json', '')
            md_filename = f"{base_name}_report.md"
            docx_filename = f"{base_name}_report.docx"
            
            md_filepath = os.path.join(markdown_folder, md_filename)
            docx_filepath = os.path.join(word_folder, docx_filename)
            
            # Step 1: Convert to Markdown
            print(f"   📝 Converting to Markdown...")
            success, json_data = convert_json_to_markdown(json_file, md_filepath)
            
            if success:
                print(f"   ✅ Markdown created: {md_filename}")
                markdown_success += 1
                
                # Step 2: Convert to Word (using python-docx)
                print(f"   📄 Converting to Word...")
                if convert_json_to_word(json_data, docx_filepath):
                    print(f"   ✅ Word created: {docx_filename}")
                    word_success += 1
                else:
                    errors += 1
            else:
                print(f"   ❌ Markdown conversion failed")
                errors += 1
            
            print()
            
        except Exception as e:
            print(f"   ❌ Error: {e}\n")
            errors += 1
    
    # Summary
    print("="*80)
    print("🎯 CONVERSION COMPLETE!")
    print("="*80)
    print(f"✅ Markdown files created: {markdown_success}")
    print(f"✅ Word files created: {word_success}")
    if errors > 0:
        print(f"❌ Errors: {errors}")
    print()
    print(f"📁 Markdown reports: {markdown_folder}")
    print(f"📁 Word reports: {word_folder}")
    print("="*80)


if __name__ == '__main__':
    # Path to the directory containing JSON files
    data_path = '/Volumes/qa_datascience_convaiqa/volumes/dataanalyticslake_convai/PRISMArchives/PI/'
    
    print("\n" + "="*80)
    print("🚀 PAGENUMBERS JSON TO MARKDOWN & WORD CONVERTER (FINAL)")
    print("="*80)
    print(f"Source: {data_path}")
    print(f"Looking for: *_WORKING_PageNumbers_analysis.json")
    print(f"Markdown Output: markdown_reports/")
    print(f"Word Output: word_reports/")
    print("="*80 + "\n")
    
    process_all_files(data_path)
    
    print("\n✅ All files converted!")
    print(f"\n📍 FINAL OUTPUT PATHS:")
    print(f"   📝 Markdown: {data_path}markdown_reports/")
    print(f"   📄 Word: {data_path}word_reports/")
    print("\n💡 Note: Uses python-docx for Word conversion (no pandoc needed!)")
    print("   Install: pip install python-docx\n")

