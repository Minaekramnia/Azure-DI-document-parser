"""
Convert PageNumbers JSON to Word Documents - FIXED for Databricks Volumes

This version handles the "Operation not supported" error by:
1. Creating Word files in /tmp/ first
2. Then copying them to the final destination

Output Path:
- Word: /Volumes/.../PI/word_reports/
"""

import json
import glob
import os
import shutil
import tempfile
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("❌ python-docx not installed. Run: pip install python-docx")
    exit(1)


def convert_json_to_word(json_filepath, output_filepath):
    """
    Convert JSON to Word document with filesystem workaround.
    
    Args:
        json_filepath: Path to JSON file
        output_filepath: Final destination for Word file
    
    Returns:
        bool: Success status
    """
    try:
        # Load JSON
        with open(json_filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Create Word document
        doc = Document()
        
        # Set default font to Cambria for the whole document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Cambria'
        font.size = Pt(11)
        
        # Title - formal, no bold/italic
        title = doc.add_paragraph('Document Analysis Report')
        title_format = title.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format.space_after = Pt(12)
        title_run = title.runs[0]
        title_run.font.name = 'Cambria'
        title_run.font.size = Pt(16)
        
        # Metadata - just date, no time, all Cambria
        p = doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        p.runs[0].font.name = 'Cambria'
        p.runs[0].font.size = Pt(11)
        
        p = doc.add_paragraph(f"Source Document: {os.path.basename(data.get('document_path', 'Unknown'))}")
        p.runs[0].font.name = 'Cambria'
        p.runs[0].font.size = Pt(11)
        
        p = doc.add_paragraph(f"Total Pages: {data.get('total_pages', 'Unknown')}")
        p.runs[0].font.name = 'Cambria'
        p.runs[0].font.size = Pt(11)
        
        p = doc.add_paragraph(f"Total Segments: {data.get('total_segments', 'Unknown')}")
        p.runs[0].font.name = 'Cambria'
        p.runs[0].font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Calculate summary statistics for the beginning
        segments = data.get('segments', [])
        total_names = sum(len(seg.get('extracted_names', [])) for seg in segments)
        total_classifications = sum(len(seg.get('classifications', [])) for seg in segments)
        
        # Category breakdown
        category_counts = {}
        for segment in segments:
            for classification in segment.get('classifications', []):
                category = classification.get('category', 'Unknown')
                category_counts[category] = category_counts.get(category, 0) + 1
        
        # Summary Statistics at the beginning
        p = doc.add_paragraph('Summary Statistics')
        p.runs[0].font.name = 'Cambria'
        p.runs[0].font.size = Pt(14)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        
        p = doc.add_paragraph(f"Total Names Extracted: {total_names}")
        p.runs[0].font.name = 'Cambria'
        p.runs[0].font.size = Pt(11)
        
        p = doc.add_paragraph(f"Total Classifications: {total_classifications}")
        p.runs[0].font.name = 'Cambria'
        p.runs[0].font.size = Pt(11)
        
        p = doc.add_paragraph(f"Total Segments: {len(segments)}")
        p.runs[0].font.name = 'Cambria'
        p.runs[0].font.size = Pt(11)
        
        p = doc.add_paragraph(f"Total Pages: {data.get('total_pages', 'Unknown')}")
        p.runs[0].font.name = 'Cambria'
        p.runs[0].font.size = Pt(11)
        
        # Classification Breakdown at the beginning
        if category_counts:
            doc.add_paragraph()
            p = doc.add_paragraph('Classification Breakdown')
            p.runs[0].font.name = 'Cambria'
            p.runs[0].font.size = Pt(12)
            p.paragraph_format.space_before = Pt(6)
            
            for category, count in sorted(category_counts.items(), key=lambda x: x[1], reverse=True):
                p = doc.add_paragraph(f"{category}: {count}")
                p.runs[0].font.name = 'Cambria'
                p.runs[0].font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Process segments
        segments = data.get('segments', [])
        
        if not segments:
            p = doc.add_paragraph("No segments found.")
            p.runs[0].font.name = 'Cambria'
            p.runs[0].font.size = Pt(11)
        
        for segment in segments:
            segment_id = segment.get('segment_id', 'Unknown')
            page_range = segment.get('page_range', 'Unknown')
            
            # Segment header - formal style, no bold
            p = doc.add_paragraph(f'Segment: {segment_id}')
            p.runs[0].font.name = 'Cambria'
            p.runs[0].font.size = Pt(14)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            p = doc.add_paragraph(f"Pages: {page_range}")
            p.runs[0].font.name = 'Cambria'
            p.runs[0].font.size = Pt(11)
            
            # Extracted Names - formal style
            p = doc.add_paragraph('Extracted Names')
            p.runs[0].font.name = 'Cambria'
            p.runs[0].font.size = Pt(12)
            p.paragraph_format.space_before = Pt(6)
            
            extracted_names = segment.get('extracted_names', [])
            if extracted_names:
                for name in extracted_names:
                    p = doc.add_paragraph(f"  - {name}")
                    p.runs[0].font.name = 'Cambria'
                    p.runs[0].font.size = Pt(11)
            else:
                p = doc.add_paragraph("No names extracted")
                p.runs[0].font.name = 'Cambria'
                p.runs[0].font.size = Pt(11)
            
            # Classifications - formal style
            p = doc.add_paragraph('Sensitive Content Classifications')
            p.runs[0].font.name = 'Cambria'
            p.runs[0].font.size = Pt(12)
            p.paragraph_format.space_before = Pt(6)
            
            classifications = segment.get('classifications', [])
            
            if not classifications:
                p = doc.add_paragraph("No sensitive content identified")
                p.runs[0].font.name = 'Cambria'
                p.runs[0].font.size = Pt(11)
            else:
                p = doc.add_paragraph(f"Total Classifications: {len(classifications)}")
                p.runs[0].font.name = 'Cambria'
                p.runs[0].font.size = Pt(11)
                doc.add_paragraph()
                
                for i, classification in enumerate(classifications, 1):
                    category = classification.get('category', 'Unknown')
                    text = classification.get('text', 'N/A')
                    page_number = classification.get('page_number', 'N/A')
                    page_range_cv = classification.get('page_range', None)
                    confidence = classification.get('confidence_score', 'N/A')
                    reason = classification.get('reason', 'N/A')
                    
                    # Classification header - no bold
                    p = doc.add_paragraph(f'Classification {i}')
                    p.runs[0].font.name = 'Cambria'
                    p.runs[0].font.size = Pt(11)
                    p.paragraph_format.space_before = Pt(6)
                    
                    # Category - no bold
                    if page_range_cv:
                        page_text = f"{page_number} (Range: {page_range_cv})"
                    else:
                        page_text = str(page_number)
                    
                    p = doc.add_paragraph(f"Category: {category}")
                    p.runs[0].font.name = 'Cambria'
                    p.runs[0].font.size = Pt(11)
                    
                    p = doc.add_paragraph(f"Page: {page_text}")
                    p.runs[0].font.name = 'Cambria'
                    p.runs[0].font.size = Pt(11)
                    
                    p = doc.add_paragraph(f"Confidence: {confidence}")
                    p.runs[0].font.name = 'Cambria'
                    p.runs[0].font.size = Pt(11)
                    
                    p = doc.add_paragraph(f"Content: {text}")
                    p.runs[0].font.name = 'Cambria'
                    p.runs[0].font.size = Pt(11)
                    
                    p = doc.add_paragraph(f"Reason: {reason}")
                    p.runs[0].font.name = 'Cambria'
                    p.runs[0].font.size = Pt(11)
                    
                    doc.add_paragraph()  # Spacing
        
        # Summary statistics are now at the beginning of the document
        
        # Footer - formal style
        doc.add_paragraph()
        footer = doc.add_paragraph("Report generated by Azure DI Output Parser")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.name = 'Cambria'
        footer.runs[0].font.size = Pt(10)
        
        # WORKAROUND: Save to temp file first, then copy
        # This avoids "Operation not supported" error on Databricks Volumes
        with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as tmp_file:
            temp_path = tmp_file.name
        
        # Save to temp location
        doc.save(temp_path)
        
        # Copy to final destination
        shutil.copy2(temp_path, output_filepath)
        
        # Clean up temp file
        try:
            os.remove(temp_path)
        except:
            pass
        
        return True
        
    except Exception as e:
        print(f"   ❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Main conversion function."""
    
    # Path
    data_path = '/Volumes/qa_datascience_convaiqa/volumes/dataanalyticslake_convai/PRISMArchives/PI/'
    
    # Find JSON files
    json_files = glob.glob(f"{data_path}*_WORKING_PageNumbers_analysis.json")
    
    if not json_files:
        print(f"❌ No JSON files found in {data_path}")
        print(f"   Looking for: *_WORKING_PageNumbers_analysis.json")
        return
    
    # Create word_reports folder
    word_folder = os.path.join(data_path, 'word_reports')
    
    try:
        os.makedirs(word_folder, exist_ok=True)
        print(f"✅ Word folder: {word_folder}\n")
    except Exception as e:
        print(f"❌ Could not create folder: {e}")
        return
    
    print("="*80)
    print("🚀 CONVERTING JSON TO WORD DOCUMENTS")
    print("="*80)
    print(f"Found {len(json_files)} JSON files")
    print(f"Output: {word_folder}")
    print("="*80 + "\n")
    
    success_count = 0
    error_count = 0
    
    for json_file in sorted(json_files):
        filename = os.path.basename(json_file)
        print(f"📄 {filename}")
        
        # Create output filename
        base_name = filename.replace('_WORKING_PageNumbers_analysis.json', '')
        word_file = os.path.join(word_folder, f"{base_name}_report.docx")
        
        # Convert
        print(f"   📄 Creating Word document...")
        if convert_json_to_word(json_file, word_file):
            print(f"   ✅ Success: {base_name}_report.docx")
            success_count += 1
        else:
            print(f"   ❌ Failed")
            error_count += 1
        
        print()
    
    # Summary
    print("="*80)
    print("🎯 CONVERSION COMPLETE!")
    print("="*80)
    print(f"✅ Word files created: {success_count}")
    if error_count > 0:
        print(f"❌ Errors: {error_count}")
    print()
    print(f"📁 Output folder: {word_folder}")
    print("="*80)
    
    print(f"\n📍 FINAL OUTPUT PATH:")
    print(f"   📄 Word documents: {word_folder}\n")


if __name__ == '__main__':
    if not DOCX_AVAILABLE:
        print("\n❌ Cannot proceed without python-docx")
        print("   Install with: pip install python-docx\n")
    else:
        main()

